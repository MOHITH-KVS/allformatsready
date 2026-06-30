import os
import io
import re
import base64
from contextlib import asynccontextmanager

from fastapi import FastAPI, File, UploadFile, HTTPException, Request
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import JSONResponse, Response
from slowapi import Limiter, _rate_limit_exceeded_handler
from slowapi.util import get_remote_address
from slowapi.errors import RateLimitExceeded

fitz_mod = None
Image_cls = None
DocxDocument = None
psd_open = None
svg2rlg = None
renderPM = None

def load_libs():
    global fitz_mod, Image_cls, DocxDocument, psd_open, svg2rlg, renderPM
    if fitz_mod is None:
        import fitz as _fitz
        fitz_mod = _fitz
    if Image_cls is None:
        from PIL import Image as _Image
        Image_cls = _Image
        # Register HEIC support
        try:
            from pillow_heif import register_heif_opener
            register_heif_opener()
        except:
            pass
    if DocxDocument is None:
        try:
            from docx import Document as _Doc
            DocxDocument = _Doc
        except:
            DocxDocument = None
    if psd_open is None:
        try:
            from psd_tools import PSDImage as _PSDImage
            psd_open = _PSDImage.open
        except:
            psd_open = False
    if svg2rlg is None:
        try:
            from svglib.svglib import svg2rlg as _svg2rlg
            from reportlab.graphics import renderPM as _renderPM
            svg2rlg = _svg2rlg
            renderPM = _renderPM
        except:
            svg2rlg = False

@asynccontextmanager
async def lifespan(app):
    load_libs()
    yield

# Rate limiter — max 10 requests per minute per IP
limiter = Limiter(key_func=get_remote_address, default_limits=["10/minute"])

app = FastAPI(title="AllFormatsReady API", lifespan=lifespan)
app.state.limiter = limiter
app.add_exception_handler(RateLimitExceeded, _rate_limit_exceeded_handler)

# CORS middleware
app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=False,
    allow_methods=["GET", "POST", "OPTIONS"],
    allow_headers=["*"],
    expose_headers=["*"],
    max_age=3600,
)

# Extra CORS safety on every response
@app.middleware("http")
async def add_cors_headers(request: Request, call_next):
    if request.method == "OPTIONS":
        return Response(
            status_code=200,
            headers={
                "Access-Control-Allow-Origin": "*",
                "Access-Control-Allow-Methods": "GET, POST, OPTIONS",
                "Access-Control-Allow-Headers": "*",
            }
        )
    response = await call_next(request)
    response.headers["Access-Control-Allow-Origin"] = "*"
    response.headers["Access-Control-Allow-Methods"] = "GET, POST, OPTIONS"
    response.headers["Access-Control-Allow-Headers"] = "*"
    return response

MAX_FILE_SIZE = 10 * 1024 * 1024
AADHAAR_RE = re.compile(r'\b\d{4}\s?\d{4}\s?\d{4}\b')
PAN_RE = re.compile(r'\b[A-Z]{5}[0-9]{4}[A-Z]\b')
SENSITIVE_KW = ["aadhaar","aadhar","uid","pan","income","passport"]


def is_sensitive(filename, text=""):
    n = filename.lower()
    if any(k in n for k in SENSITIVE_KW): return True
    if text and (AADHAAR_RE.search(text) or PAN_RE.search(text)): return True
    return False


def pdf_page_to_pil(page, dpi=120):
    mat = fitz_mod.Matrix(dpi/72, dpi/72)
    pix = page.get_pixmap(matrix=mat, alpha=False)
    return Image_cls.frombytes("RGB", [pix.width, pix.height], pix.samples)


def compress_to_target(img, fmt, target_kb):
    target_bytes = target_kb * 1024
    lo, hi, best = 10, 95, None
    for _ in range(10):
        mid = (lo + hi) // 2
        buf = io.BytesIO()
        if fmt == "JPEG": img.save(buf, format="JPEG", quality=mid, optimize=True)
        elif fmt == "WEBP": img.save(buf, format="WEBP", quality=mid)
        size = buf.tell()
        if size <= target_bytes: best = buf.getvalue(); lo = mid + 1
        else: hi = mid - 1
    if best is None:
        buf = io.BytesIO()
        if fmt == "JPEG": img.save(buf, format="JPEG", quality=10, optimize=True)
        elif fmt == "WEBP": img.save(buf, format="WEBP", quality=10)
        best = buf.getvalue()
    return best


def compress_pdf(pdf_bytes):
    doc = fitz_mod.open(stream=pdf_bytes, filetype="pdf")
    buf = io.BytesIO()
    doc.save(buf, garbage=4, deflate=True, deflate_images=True, deflate_fonts=True, clean=True)
    doc.close()
    return buf.getvalue()


def mask_pdf(pdf_bytes):
    doc = fitz_mod.open(stream=pdf_bytes, filetype="pdf")
    for page in doc:
        for block in page.get_text("dict")["blocks"]:
            if block.get("type") != 0: continue
            for line in block.get("lines", []):
                for span in line.get("spans", []):
                    for m in AADHAAR_RE.finditer(span["text"]):
                        for rect in page.search_for(m.group(0)):
                            r = fitz_mod.Rect(rect.x0, rect.y0, rect.x0 + rect.width*0.67, rect.y1)
                            page.add_redact_annot(r, fill=(0,0,0))
        page.apply_redactions()
    buf = io.BytesIO()
    doc.save(buf, garbage=4, deflate=True)
    doc.close()
    return buf.getvalue()


def img_bytes(img, fmt, quality=85):
    buf = io.BytesIO()
    if fmt == "JPEG": img.save(buf, format="JPEG", quality=quality, optimize=True)
    elif fmt == "PNG": img.save(buf, format="PNG", optimize=True)
    elif fmt == "WEBP": img.save(buf, format="WEBP", quality=quality)
    elif fmt == "PDF": img.save(buf, format="PDF")
    return buf.getvalue()


def pil_to_docx(imgs) -> bytes:
    """Embed one or more PIL images into a single DOCX file — one image per page."""
    if DocxDocument is None:
        return b""
    try:
        doc = DocxDocument()
        # If single image passed, wrap in list
        if not isinstance(imgs, list):
            imgs = [imgs]
        for idx, img in enumerate(imgs):
            buf = io.BytesIO()
            img.save(buf, format="PNG")
            buf.seek(0)
            doc.add_picture(buf, width=None)
            # Add page break between pages (not after last)
            if idx < len(imgs) - 1:
                from docx.oxml.ns import qn
                from docx.oxml import OxmlElement
                p = doc.add_paragraph()
                run = p.add_run()
                br = OxmlElement('w:br')
                br.set(qn('w:type'), 'page')
                run._r.append(br)
        out = io.BytesIO()
        doc.save(out)
        return out.getvalue()
    except:
        return b""


def psd_to_pil(file_bytes):
    """Convert PSD bytes to a flattened PIL RGB image."""
    psd = psd_open(io.BytesIO(file_bytes))
    img = psd.composite()  # flattens all visible layers
    if img.mode != "RGB":
        img = img.convert("RGB")
    return img


def svg_to_pil(file_bytes):
    """Convert SVG bytes to a PIL RGB image via svglib + reportlab."""
    import tempfile
    with tempfile.NamedTemporaryFile(suffix=".svg", delete=False) as tmp:
        tmp.write(file_bytes)
        tmp_path = tmp.name
    try:
        drawing = svg2rlg(tmp_path)
        png_buf = io.BytesIO()
        renderPM.drawToFile(drawing, png_buf, fmt="PNG")
        png_buf.seek(0)
        img = Image_cls.open(png_buf)
        if img.mode in ("RGBA", "P"):
            bg = Image_cls.new("RGB", img.size, (255, 255, 255))
            bg.paste(img.convert("RGBA"), mask=img.convert("RGBA").split()[-1])
            img = bg
        else:
            img = img.convert("RGB")
        return img
    finally:
        try:
            os.unlink(tmp_path)
        except:
            pass


def open_any_image(file_bytes, ext, content_type):
    """Open an image of any supported format (incl. GIF/BMP/TIFF/PSD/SVG) as a flattened PIL RGB image."""
    is_psd = ext == "psd" or "photoshop" in (content_type or "")
    is_svg = ext == "svg" or "svg" in (content_type or "")

    if is_psd:
        if not psd_open:
            raise HTTPException(status_code=400, detail="PSD support is currently unavailable.")
        return psd_to_pil(file_bytes)

    if is_svg:
        if not svg2rlg:
            raise HTTPException(status_code=400, detail="SVG support is currently unavailable.")
        return svg_to_pil(file_bytes)

    # GIF / BMP / TIFF / JPG / PNG / WEBP / HEIC all handled natively by Pillow
    img = Image_cls.open(io.BytesIO(file_bytes))
    if getattr(img, "is_animated", False):
        img.seek(0)  # use first frame for animated GIFs
    if img.mode != "RGB":
        img = img.convert("RGB")
    return img


def make_file(name, fmt, data, label, category, page=None):
    return {
        "name": name, "format": fmt, "label": label,
        "category": category, "page": page,
        "size_bytes": len(data), "size_kb": round(len(data)/1024, 1),
        "data_b64": base64.b64encode(data).decode(),
    }


def image_outputs(img, prefix="", page_num=None):
    """Generate all image formats for a single page — NO DOCX here, handled separately."""
    outputs = []
    pg = f" · Page {page_num}" if page_num else ""
    p = f"p{page_num}_" if page_num else ""

    # Resize if too large
    max_w = 1200
    if img.width > max_w:
        ratio = max_w / img.width
        img = img.resize((max_w, int(img.height * ratio)), Image_cls.LANCZOS)

    # JPG variants
    for q, lbl in [(85,"High Quality"),(50,"Compressed")]:
        d = img_bytes(img, "JPEG", q)
        outputs.append(make_file(f"{p}jpg_{lbl.lower().replace(' ','_')}.jpg","JPG",d,f"JPG — {lbl}{pg}","JPG",page_num))

    # JPG size targets
    for kb in [100, 200, 500]:
        d = compress_to_target(img, "JPEG", kb)
        outputs.append(make_file(f"{p}jpg_under_{kb}kb.jpg","JPG",d,f"JPG — Under {kb}KB{pg}","JPG",page_num))

    # PNG
    d = img_bytes(img, "PNG")
    outputs.append(make_file(f"{p}image.png","PNG",d,f"PNG — Lossless{pg}","PNG",page_num))

    # WebP
    d = compress_to_target(img, "WEBP", 200)
    outputs.append(make_file(f"{p}webp_under_200kb.webp","WEBP",d,f"WebP — Under 200KB{pg}","WebP",page_num))

    return outputs


@app.get("/")
def root(): return {"status": "AllFormatsReady API is live"}

@app.get("/health")
def health(): return {"status": "ok"}

@app.get("/ping")
def ping(): return {"ping": "pong"}


@app.post("/convert")
@limiter.limit("10/minute")
async def convert(request: Request, file: UploadFile = File(...)):
    load_libs()
    filename = file.filename or "document"
    ext = filename.lower().rsplit(".", 1)[-1] if "." in filename else ""
    file_bytes = await file.read()

    if len(file_bytes) > MAX_FILE_SIZE:
        raise HTTPException(status_code=400, detail="File too large. Max 10MB allowed.")

    outputs = []
    total_pages = 1

    is_pdf = ext == "pdf" or "pdf" in (file.content_type or "")
    is_docx = ext in ["docx","doc"] or "wordprocessing" in (file.content_type or "")
    is_image = not is_pdf and not is_docx

    if is_docx:
        raise HTTPException(
            status_code=400,
            detail="DOCX upload not supported. Please convert to PDF first, then upload."
        )

    elif is_pdf:
        doc = fitz_mod.open(stream=file_bytes, filetype="pdf")
        total_pages = len(doc)
        text = "".join(pg.get_text() for pg in doc)
        sensitive = is_sensitive(filename, text)

        outputs.append(make_file("compressed.pdf","PDF",compress_pdf(file_bytes),"Compressed PDF","PDF"))

        if sensitive:
            try:
                outputs.append(make_file("masked_aadhaar.pdf","PDF",mask_pdf(file_bytes),"Masked Aadhaar PDF","PDF"))
            except: pass

        all_page_imgs = []
        for i, page in enumerate(doc, 1):
            img = pdf_page_to_pil(page)
            max_w = 1200
            if img.width > max_w:
                ratio = max_w / img.width
                img = img.resize((max_w, int(img.height * ratio)), Image_cls.LANCZOS)
            all_page_imgs.append(img)

        doc.close()

        for i, img in enumerate(all_page_imgs, 1):
            pg = i if total_pages > 1 else None
            outputs.extend(image_outputs(img, f"page{i}", pg))

        d = pil_to_docx(all_page_imgs)
        if d:
            label = f"DOCX — All {total_pages} Pages" if total_pages > 1 else "DOCX — Word Document"
            outputs.append(make_file("document_all_pages.docx","DOCX",d,label,"DOCX"))

    elif is_image:
        try:
            img = open_any_image(file_bytes, ext, file.content_type)
        except HTTPException:
            raise
        except:
            raise HTTPException(status_code=400, detail="Could not read image file.")

        outputs.extend(image_outputs(img))

        raw_pdf = img_bytes(img, "PDF")
        outputs.append(make_file("image_as_pdf.pdf","PDF",raw_pdf,"PDF — From Image","PDF"))
        outputs.append(make_file("image_as_pdf_compressed.pdf","PDF",compress_pdf(raw_pdf),"PDF — Compressed","PDF"))

        d = pil_to_docx([img])
        if d:
            outputs.append(make_file("document.docx","DOCX",d,"DOCX — Word Document","DOCX"))

    return JSONResponse(content={"files": outputs, "total": len(outputs), "total_pages": total_pages})


@app.post("/convert-multiple")
@limiter.limit("5/minute")
async def convert_multiple(request: Request, files: list[UploadFile] = File(...)):
    """Accept multiple images → generate individual formats + one combined PDF + combined DOCX."""
    load_libs()

    if len(files) > 10:
        raise HTTPException(status_code=400, detail="Maximum 10 images at once.")

    all_imgs = []
    total_size = 0

    for f in files:
        file_bytes = await f.read()
        total_size += len(file_bytes)
        if total_size > MAX_FILE_SIZE:
            raise HTTPException(status_code=400, detail="Total size too large. Max 10MB total.")
        fname = f.filename or ""
        f_ext = fname.lower().rsplit(".", 1)[-1] if "." in fname else ""
        try:
            img = open_any_image(file_bytes, f_ext, f.content_type)
            # Resize if needed
            max_w = 1200
            if img.width > max_w:
                ratio = max_w / img.width
                img = img.resize((max_w, int(img.height * ratio)), Image_cls.LANCZOS)
            all_imgs.append((fname or f"image_{len(all_imgs)+1}", img))
        except HTTPException:
            raise
        except:
            raise HTTPException(status_code=400, detail=f"Could not read image: {f.filename}")

    if not all_imgs:
        raise HTTPException(status_code=400, detail="No valid images found.")

    outputs = []
    total_count = len(all_imgs)

    # ── Per-image individual outputs ──
    for i, (fname, img) in enumerate(all_imgs, 1):
        pg = i if total_count > 1 else None
        pg_label = f" · Image {i}" if total_count > 1 else ""
        p = f"img{i}_" if total_count > 1 else ""

        # JPG variants
        for q, lbl in [(85,"High Quality"),(50,"Compressed")]:
            d = img_bytes(img, "JPEG", q)
            outputs.append(make_file(f"{p}jpg_{lbl.lower().replace(' ','_')}.jpg","JPG",d,f"JPG — {lbl}{pg_label}","JPG",pg))

        # JPG size targets
        for kb in [100, 200, 500]:
            d = compress_to_target(img, "JPEG", kb)
            outputs.append(make_file(f"{p}jpg_under_{kb}kb.jpg","JPG",d,f"JPG — Under {kb}KB{pg_label}","JPG",pg))

        # PNG
        d = img_bytes(img, "PNG")
        outputs.append(make_file(f"{p}image.png","PNG",d,f"PNG — Lossless{pg_label}","PNG",pg))

        # WebP
        d = compress_to_target(img, "WEBP", 200)
        outputs.append(make_file(f"{p}webp_under_200kb.webp","WEBP",d,f"WebP — Under 200KB{pg_label}","WebP",pg))

        # Individual PDF
        raw_pdf = img_bytes(img, "PDF")
        outputs.append(make_file(f"{p}image_as_pdf.pdf","PDF",raw_pdf,f"PDF — Image {i}","PDF",pg))

    # ── Combined PDF (all images in one PDF) ──
    if total_count > 1:
        imgs_only = [img for _, img in all_imgs]
        # Use Pillow to save all images as multi-page PDF
        buf = io.BytesIO()
        imgs_only[0].save(buf, format="PDF", save_all=True, append_images=imgs_only[1:])
        combined_pdf = buf.getvalue()
        outputs.append(make_file(
            "combined_all_images.pdf","PDF",
            compress_pdf(combined_pdf),
            f"PDF — All {total_count} Images Combined","PDF"
        ))

        # ── Combined DOCX (all images in one Word doc) ──
        d = pil_to_docx(imgs_only)
        if d:
            outputs.append(make_file(
                "combined_all_images.docx","DOCX",d,
                f"DOCX — All {total_count} Images Combined","DOCX"
            ))

    return JSONResponse(content={
        "files": outputs,
        "total": len(outputs),
        "total_pages": total_count,
        "combined": total_count > 1
    })
